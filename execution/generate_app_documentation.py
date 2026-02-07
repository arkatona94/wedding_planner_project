"""
Generate Application Documentation Word Document
Creates a comprehensive Word document outlining the EverAfter Wedding Planner application.
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_documentation():
    """Create the application documentation Word document."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('EverAfter Wedding Planner', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Application Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Application Overview',
        '2. Technology Stack',
        '3. Database Architecture (Supabase)',
        '4. Vendor Management System',
        '5. API Keys & Integrations',
        '6. Key Features',
        '7. File Structure'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Application Overview
    doc.add_heading('1. Application Overview', level=1)
    doc.add_paragraph(
        'EverAfter is a comprehensive wedding planning "Command Center" that combines the best features '
        'from top wedding apps (The Knot, Zola, Joy, Bridebook, Prismm) into a single, privacy-first application. '
        'The application provides couples with all the tools they need to plan their perfect wedding day.'
    )
    
    doc.add_heading('Key Capabilities', level=2)
    capabilities = [
        ('Planning Hub', 'Adaptive checklists, budget tracker, visual dashboard with charts and countdown'),
        ('Guest Management', 'RSVP tracking, CSV import/export, QR code RSVP functionality'),
        ('Visual Tools', 'Drag-and-drop seating charts, timeline builder, customizable wedding website'),
        ('Vendor Management', 'Track contacts, contracts, pricing, ratings, deposits and payments'),
        ('Photo Gallery', 'Zero-app photo sharing via QR code - no app download required for guests'),
        ('AI Features', 'Virtual dress try-on via Replicate, AI-powered recommendations via Gemini'),
        ('Communications', 'Email and SMS notifications to guests via Resend and Twilio')
    ]
    for name, desc in capabilities:
        doc.add_paragraph(f'{name}: {desc}', style='List Bullet')
    
    doc.add_page_break()
    
    # 2. Technology Stack
    doc.add_heading('2. Technology Stack', level=1)
    
    doc.add_heading('Frontend', level=2)
    frontend_tech = [
        ('React 18', 'Core UI framework with TypeScript for type safety'),
        ('Vite', 'Build tool for fast development and optimized production builds'),
        ('Tailwind CSS', 'Utility-first CSS framework for styling'),
        ('Zustand', 'Lightweight state management'),
        ('Recharts', 'Chart library for budget visualizations'),
        ('React Router v6', 'Client-side routing')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Purpose'
    for tech, purpose in frontend_tech:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose
    
    doc.add_paragraph()
    
    doc.add_heading('Backend', level=2)
    backend_tech = [
        ('Supabase', 'PostgreSQL database, authentication, storage, and edge functions'),
        ('Python Scripts', 'Data processing, web scraping, communication utilities'),
        ('Deno', 'Runtime for Supabase Edge Functions')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Purpose'
    for tech, purpose in backend_tech:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose
    
    doc.add_page_break()
    
    # 3. Database Architecture
    doc.add_heading('3. Database Architecture (Supabase)', level=1)
    
    doc.add_heading('Database Details', level=2)
    db_info = [
        ('Provider', 'Supabase (PostgreSQL)'),
        ('Project ID', 'xcjelqmifskowxxdtqrh'),
        ('URL', 'https://xcjelqmifskowxxdtqrh.supabase.co'),
        ('Security', 'Row Level Security (RLS) enabled on all tables')
    ]
    for key, value in db_info:
        doc.add_paragraph(f'{key}: {value}', style='List Bullet')
    
    doc.add_heading('Database Tables', level=2)
    
    tables_info = [
        ('profiles', 'User profile information including email, name, avatar, and app settings'),
        ('weddings', 'Wedding event details - partner names, date, budget, guest count, venue'),
        ('checklist_items', 'Task management - title, description, category, due date, priority, status'),
        ('budget_items', 'Budget tracking - category, estimated vs actual costs, payments, due dates'),
        ('guests', 'Guest list - contact info, RSVP status, meal choice, dietary restrictions, seating'),
        ('vendors', 'Vendor management - name, category, contact info, price, rating, contract status'),
        ('timeline_events', 'Wedding day schedule - time, title, description, location'),
        ('seating_tables', 'Table layout - name, capacity, shape, position, rotation'),
        ('room_elements', 'Venue decor elements - type, position, dimensions, rotation'),
        ('photos', 'Photo gallery - URL, caption, category, favorite status'),
        ('communication_logs', 'Email/SMS history - recipient, channel, status, timestamps')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Description'
    for tbl_name, desc in tables_info:
        row_cells = table.add_row().cells
        row_cells[0].text = tbl_name
        row_cells[1].text = desc
    
    doc.add_paragraph()
    doc.add_heading('Database Triggers', level=2)
    doc.add_paragraph(
        'The database includes automatic triggers for:\n'
        '• handle_new_user: Creates profile and initial wedding record on user signup\n'
        '• update_updated_at_column: Automatically updates timestamps on record changes'
    )
    
    doc.add_page_break()
    
    # 4. Vendor Management System
    doc.add_heading('4. Vendor Management System', level=1)
    
    doc.add_paragraph(
        'The application includes a comprehensive vendor management system with location-aware pricing '
        'and pre-populated vendor data for the Cincinnati, OH area.'
    )
    
    doc.add_heading('Vendor Categories', level=2)
    vendor_categories = [
        ('Venue', 'Ceremony and reception locations'),
        ('Catering', 'Food and beverage services'),
        ('Photography', 'Professional photographers'),
        ('Videography', 'Wedding video production'),
        ('Florist', 'Flowers and floral arrangements'),
        ('Music/DJ', 'DJs, live bands, musicians'),
        ('Officiant', 'Wedding ceremony officiants'),
        ('Cake', 'Wedding cakes and desserts'),
        ('Rentals', 'Equipment and furniture rentals'),
        ('Transportation', 'Limos, shuttles, classic cars'),
        ('Hair & Makeup', 'Bridal beauty services')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Description'
    for cat, desc in vendor_categories:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('Location-Aware Features', level=2)
    doc.add_paragraph(
        'The vendor generator and budget calculator use regional cost multipliers based on:\n'
        '• State-level cost-of-living data (50 states covered)\n'
        '• City-specific adjustments for major metros\n'
        '• Data sourced from The Knot 2024 Real Weddings Study'
    )
    
    doc.add_heading('Sample Regional Multipliers', level=3)
    multipliers = [
        ('New York, NY', '2.00x (Highest)'),
        ('San Francisco, CA', '1.85x'),
        ('Los Angeles, CA', '1.70x'),
        ('Chicago, IL', '1.35x'),
        ('Cincinnati, OH', '0.90x (Baseline data)'),
        ('Indianapolis, IN', '0.80x'),
        ('Mississippi', '0.70x (Lowest)')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Location'
    hdr_cells[1].text = 'Cost Multiplier'
    for loc, mult in multipliers:
        row_cells = table.add_row().cells
        row_cells[0].text = loc
        row_cells[1].text = mult
    
    doc.add_page_break()
    
    # 5. API Keys & Integrations
    doc.add_heading('5. API Keys & Integrations', level=1)
    
    doc.add_paragraph(
        'The application integrates with several third-party services. All API keys are stored '
        'in the .env file at the project root and should never be committed to version control.'
    )
    
    doc.add_heading('Integration Summary', level=2)
    
    integrations = [
        ('Supabase', 'Database & Authentication', 
         'PostgreSQL database, user authentication, file storage, edge functions',
         'VITE_SUPABASE_URL, VITE_SUPABASE_ANON_KEY, SUPABASE_SERVICE_ROLE_KEY'),
        ('Google Gemini AI', 'AI Features',
         'AI-powered recommendations, content generation, smart suggestions',
         'VITE_GEMINI_API_KEY'),
        ('Replicate', 'Virtual Try-On',
         'IDM-VTON model for virtual wedding dress try-on feature',
         'VITE_REPLICATE_API_TOKEN'),
        ('Twilio', 'SMS Notifications',
         'Send SMS messages to guests for RSVP reminders and updates',
         'TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN, TWILIO_PHONE_NUMBER'),
        ('Resend', 'Email Notifications',
         'Send transactional emails for RSVP, reminders, updates',
         'RESEND_API_KEY, RESEND_FROM_EMAIL'),
        ('GitHub', 'Version Control',
         'Repository management and collaboration',
         'GITHUB_PERSONAL_ACCESS_TOKEN')
    ]
    
    for svc, category, desc, keys in integrations:
        doc.add_heading(f'{svc} ({category})', level=3)
        doc.add_paragraph(f'Purpose: {desc}')
        doc.add_paragraph(f'Environment Variables: {keys}')
    
    doc.add_heading('API Key Status', level=2)
    key_status = [
        ('Supabase', 'Configured', 'Project: xcjelqmifskowxxdtqrh'),
        ('Google Gemini', 'Configured', 'Active API key'),
        ('Replicate', 'Configured', 'IDM-VTON v1.5 model'),
        ('Twilio', 'Configured', 'Account SID and Auth Token set'),
        ('Resend', 'Placeholder', 'Needs real API key for production')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Service'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Notes'
    for svc, status, notes in key_status:
        row_cells = table.add_row().cells
        row_cells[0].text = svc
        row_cells[1].text = status
        row_cells[2].text = notes
    
    doc.add_page_break()
    
    # 6. Key Features
    doc.add_heading('6. Key Features', level=1)
    
    features = [
        ('Budget Calculator', 
         'Location-aware budget allocation based on industry standards. Uses The Knot 2024 data '
         'with regional cost adjustments. Provides recommended budget breakdowns across 15 categories.'),
        ('Guest Management',
         'Full RSVP tracking with meal choices, dietary restrictions, plus-ones, and table assignments. '
         'Supports CSV import/export and QR code RSVP.'),
        ('Seating Charts',
         'Drag-and-drop interface for creating table layouts. Supports round, rectangular, and custom tables. '
         'Includes room elements like dance floor, stage, bar, etc.'),
        ('Timeline Builder',
         'Create detailed day-of schedules with auto-adjustment when event times change. '
         'Supports multiple day types (rehearsal dinner, wedding day, etc.)'),
        ('Vendor Tracking',
         'Track vendor contacts, contracts, pricing, deposits, and ratings. '
         'Location-aware vendor suggestions with pricing estimates.'),
        ('Virtual Dress Try-On',
         'AI-powered virtual try-on using Replicate IDM-VTON model. '
         'Upload bride photo and dress image to see virtual combination.'),
        ('Communication Center',
         'Send email and SMS notifications to guests directly from the app. '
         'Logs all communications for tracking.')
    ]
    
    for name, desc in features:
        doc.add_heading(name, level=2)
        doc.add_paragraph(desc)
    
    doc.add_page_break()
    
    # 7. File Structure
    doc.add_heading('7. File Structure', level=1)
    
    structure = """
Project Root
├── src/                    # React frontend source
│   ├── components/         # Reusable UI components
│   ├── pages/              # Page components (20 pages)
│   ├── store/              # Zustand state management
│   ├── types/              # TypeScript type definitions
│   ├── utils/              # Utility functions
│   │   ├── budgetCalculator.ts    # Budget allocation logic
│   │   └── vendorGenerator.ts     # Location-aware vendor generation
│   ├── data/               # Static JSON data (vendor lists)
│   └── lib/                # External library configs
├── supabase/               # Supabase configuration
│   ├── functions/          # Edge functions (Deno)
│   │   └── send-notification/  # Email/SMS sending
│   └── migrations/         # Database migrations
├── execution/              # Python scripts
│   ├── venue_scraper.py    # Web scraping for venues
│   ├── send_communication.py  # Communication utilities
│   └── test_*.py           # Test scripts
├── directives/             # SOPs and documentation
├── backend/                # Backend utilities
├── .env                    # Environment variables (not in git)
├── package.json            # Node.js dependencies
├── tailwind.config.js      # Tailwind CSS configuration
├── vite.config.ts          # Vite build configuration
└── README.md               # Project documentation
"""
    
    # Add as preformatted text
    para = doc.add_paragraph()
    run = para.add_run(structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # Final page
    doc.add_page_break()
    doc.add_heading('Document Information', level=1)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    doc.add_paragraph('Application: EverAfter Wedding Planner')
    doc.add_paragraph('Version: 1.0')
    doc.add_paragraph('Author: AI Assistant (Antigravity)')
    
    # Save document
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 
                                'EverAfter_Application_Documentation.docx')
    doc.save(output_path)
    print(f"✅ Documentation saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    create_documentation()
